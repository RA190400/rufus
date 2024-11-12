import time
import json
import numpy as np
from collections import deque
from transformers import pipeline, AutoModel, AutoTokenizer
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
import validators
import faiss
from docx import Document
import torch
import logging
import requests
from sklearn.metrics.pairwise import cosine_similarity

instruction_message = """
You are an intelligent assistant. Please extract the answer to the question from the provided context.
If you cannot find a relevant answer in the text, return "None" or "No answer found".
"""

class Rufus:
    def __init__(self, max_input_length=450, similarity_threshold=0.3):
        self.embedding_model_name = "sentence-transformers/all-MiniLM-L6-v2"
        self.tokenizer = AutoTokenizer.from_pretrained(self.embedding_model_name)
        self.model = AutoModel.from_pretrained(self.embedding_model_name)
        
        self.embedding_dimension = 384  # Adjust based on the embedding model
        self.index = faiss.IndexFlatL2(self.embedding_dimension)
        self.texts = []
        
        self.MAX_INPUT_LENGTH = max_input_length
        self.SIMILARITY_THRESHOLD = similarity_threshold

    def get_embedding(self, text):
        inputs = self.tokenizer(text, return_tensors="pt", padding=True, truncation=True)
        inputs = inputs.to('cuda' if torch.cuda.is_available() else 'cpu')
        with torch.no_grad():
            embeddings = self.model(**inputs).last_hidden_state.mean(dim=1).cpu().numpy()
        return embeddings

    def store_chunks_in_faiss(self, chunks):
        for chunk in chunks:
            embedding = self.get_embedding(chunk)
            self.index.add(embedding)
            self.texts.append(chunk)

    def filter_relevant_chunks(self, question_embedding, chunks, threshold=0.3):
        relevant_chunks = []
        for chunk in chunks:
            chunk_embedding = self.get_embedding(chunk)
        
            # Flatten embeddings if they are 3D to 2D
            question_embedding = question_embedding.flatten() if question_embedding.ndim > 1 else question_embedding
            chunk_embedding = chunk_embedding.flatten() if chunk_embedding.ndim > 1 else chunk_embedding
        
            similarity = cosine_similarity([question_embedding], [chunk_embedding])[0][0]
            if similarity >= threshold:
                relevant_chunks.append(chunk)
    
        return relevant_chunks

    def retrieve_relevant_chunks(self, question, top_k=7):
        query_embedding = self.get_embedding(question)
        chunks_to_check = [self.texts[i] for i in range(self.index.ntotal)]  # Assume all stored chunks
        filtered_chunks = self.filter_relevant_chunks(query_embedding, chunks_to_check, self.SIMILARITY_THRESHOLD)
        
        if filtered_chunks:
            distances, indices = self.index.search(query_embedding, top_k)
            return [self.texts[i] for i in indices[0] if self.texts[i] in filtered_chunks]
        else:
            print("No relevant chunks found.")
            return None

    def get_all_relevant_urls(self, start_url, max_depth=2,max_urls=30):
        visited = set()
        queue = deque([(start_url, 0)])
        all_urls = []

        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=chrome_options)

        while queue and len(all_urls) < max_urls:
            current_url, depth = queue.popleft()
            if depth > max_depth or current_url in visited:
                continue

            visited.add(current_url)

            try:
                driver.get(current_url)
                time.sleep(2)  # Adjust sleep time as necessary
                all_urls.append(current_url)

                links = driver.find_elements(By.TAG_NAME, 'a')
                for link in links:
                    href = link.get_attribute('href')
                    if href and validators.url(href) and href not in visited:
                        queue.append((href, depth + 1))
                        if len(all_urls) >= max_urls:
                            break
            except Exception as e:
                print(f"An error occurred while accessing {current_url}: {e}")

        driver.quit()
        return all_urls

    def extract_data_from_url(self, url):
        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=chrome_options)

        try:
            driver.get(url)
            time.sleep(2)  # Adjust sleep time as necessary
            body = driver.find_element(By.TAG_NAME, 'body')
            full_text = body.text
            return full_text
        except Exception as e:
            print(f"An error occurred while accessing {url}: {e}")
            return None
        finally:
            driver.quit()

    def chunk_text(self, text, max_length):
        sentences = text.split('. ')
        chunks = []
        current_chunk = []

        for sentence in sentences:
            if len(' '.join(current_chunk)) + len(sentence) + 1 <= max_length:
                current_chunk.append(sentence)
            else:
                chunks.append('. '.join(current_chunk) + '.')
                current_chunk = [sentence]

        if current_chunk:
            chunks.append('. '.join(current_chunk) + '.')

        return chunks

    def ollama_extract_text_or_json(self, text, user_prompt):
        messages = [
            {
                "role": "user",
                "content": f"{instruction_message}\n\n{text}\n\n{user_prompt}"
            }
        ]
        payload = {
            "model": "mistral",
            "messages": messages,
            "stream": False
        }
        url = "http://127.0.0.1:11434/api/chat"

        try:
            response = requests.post(url, json=payload)
            response.raise_for_status()

            structured_data = response.json().get("message", {}).get("content", "").strip()
            logging.info(f"Raw Structured Data: {structured_data}")

            try:
                return json.loads(structured_data)
            except json.JSONDecodeError:
                if structured_data.lower() in ["none", "no answer found"]:
                    return None
                logging.warning("Response is not valid JSON. Returning structured text.")
                return structured_data

        except requests.RequestException as e:
            logging.error(f"Request error: {e}")
            return None

    def save_to_json(self, data, filename):
        with open(filename, 'w') as json_file:
            json.dump(data, json_file, indent=4)
        print(f"Data saved to {filename}")

    def save_to_docx(self, data, filename):
        doc = Document()
        doc.add_heading('RAG System Report', level=1)

        for entry in data:
            doc.add_heading(entry['question'], level=2)
            doc.add_paragraph(f"Extracted Context: {entry['context']}")
            doc.add_paragraph(f"Generated Answer: {entry['answer']}\n")
        
        doc.save(filename)
        print(f"Data saved to {filename}")

    def run(self, target_url, question, file_format):
        main_url_text = self.extract_data_from_url(target_url)
        relevant_data_found = False
        self.texts.clear()
        self.index.reset()  # Reset the FAISS index
        results = []

        if main_url_text:
            chunks = self.chunk_text(main_url_text, self.MAX_INPUT_LENGTH)
            self.store_chunks_in_faiss(chunks)

            relevant_chunks = self.retrieve_relevant_chunks(question)
            if relevant_chunks:  # Only proceed if relevant chunks are found
                combined_context = ' '.join(relevant_chunks)
                long_answer = self.ollama_extract_text_or_json(combined_context, question)
                
                if long_answer is not None:
                    print(f"Generated Answer from main URL ({target_url}):\n{long_answer}\n")
                    results.append({
                        'question': question,
                        'context': combined_context,
                        'answer': long_answer
                    })
                    relevant_data_found = True
            else:
                print("No relevant chunks found in the main URL.")

        if not relevant_data_found:
            relevant_urls = self.get_all_relevant_urls(target_url)
            for url in relevant_urls:
                print(f"Checking relevant URL: {url}")
                text = self.extract_data_from_url(url)
                if not text:
                    print(f"No text found at {url}, moving to next URL.")
                    continue

                chunks = self.chunk_text(text, self.MAX_INPUT_LENGTH)
                self.store_chunks_in_faiss(chunks)
                relevant_chunks = self.retrieve_relevant_chunks(question)

                if not relevant_chunks:  # If no relevant chunks are found, skip to next URL
                    print(f"No relevant chunks found at {url}, moving to next URL.")
                    continue

                combined_context = ' '.join(relevant_chunks)
                long_answer = self.ollama_extract_text_or_json(combined_context, question)
                
                if long_answer is not None:
                    print(f"Generated Long Answer from {url}:\n{long_answer}\n")
                    results.append({
                        'question': question,
                        'context': combined_context,
                        'answer': long_answer
                    })
                    relevant_data_found = True
                    break
                else:
                    print(f"No relevant data found on {url}, moving to next URL.")

        if not relevant_data_found:
            print("No relevant data found in the provided URLs.")
            results.append({
                    'question': question,
                    'context': "No relevant context found.",
                    'answer': "No answer found in all the targetr url and all other urls."
                })

        
        if results:
            if file_format == 'json':
                self.save_to_json(results, 'rag_results.json')
            elif file_format == 'docx':
                self.save_to_docx(results, 'rag_results.docx')
            else:
                print("Invalid file format specified.")



